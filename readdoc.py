import os
from docx import Document
import xlwt

if __name__ == '__main__':
    name_list = []
    lat_list = []
    lng_list = []
    content_list = []
    count = 1
    is_content = False
    also_content = False
    path = os.getcwd() + r'\\' + "中国海域海岛地名志（辽宁海岛卷）.docx"
    document = Document(path)
    content = ''
    print(len(document.paragraphs))
    for paragraph in document.paragraphs:
        # 获取段落的 run 对象列表
        runs = paragraph.runs
        # 获取 run 对象
        run_0 = runs[0]
        print(count)
        # print(run_0.text)
        if run_0.font.bold is None:
            print(paragraph.text)
            if is_content:
                also_content = True
            is_content = True
            if len(content_list) and also_content:
                content = content_list.pop(-1)
            content = content + paragraph.text
            content_list.append(content)
            if "北纬" in paragraph.text:
                print(paragraph.text.split('。')[0])
                N = paragraph.text.split('。')[0].split('，')[0]
                E = paragraph.text.split('。')[0].split('，')[1]
                N1 = N.split('纬')[1].split('°')[0]
                lat = float(N1)
                lat = lat + float(N.split('纬')[1].split('°')[1][:-1]) / 60
                E1 = E.split('经')[1].split('°')[0]
                lng = float(E1)
                lng = lng + float(E.split('经')[1].split('°')[1][:-1]) / 60
                # loc = N + 'N,' + E + 'E'
                # loc_list.append(loc)
                lat_list.append(lat)
                lng_list.append(lng)
        else:
            if is_content:
                content = ''
                is_content = False
                also_content = False
            name = paragraph.text  # .split('（')[0]
            print(name)
            name_list.append(name)
        count += 1
    print(len(lng_list))
    print(len(lat_list))
    print(len(content_list))
    print(len(name_list))
    print('是否加粗：', run_0.font.bold)
    print(paragraph.text)
    print(len(name_list))
    # print(len(loc_list))
    workbook = xlwt.Workbook()
    # 创建sheet对象，新建sheet
    sheet1 = workbook.add_sheet('海岛数据', cell_overwrite_ok=True)
    # ---设置excel样式---
    # 初始化样式
    style = xlwt.XFStyle()
    # 创建字体样式
    font = xlwt.Font()
    font.name = '宋体'
    #    font.bold = True #加粗
    # 设置字体
    style.font = font
    # 使用样式写入数据
    print('正在存储数据，请勿打开excel')
    # 向sheet中写入数据
    title_list = ['海岛名称', '内容', '纬度', '经度']
    for cc in range(0, len(title_list)):
        sheet1.write(0, cc, title_list[cc], style)
    for i in range(0, len(name_list)):
        print(name_list[i])
        sheet1.write(i + 1, 0, name_list[i], style)  # 海岛名称
        sheet1.write(i + 1, 1, content_list[i], style)  # 内容
        sheet1.write(i + 1, 2, lng_list[i], style)  # 纬度
        sheet1.write(i + 1, 3, lat_list[i], style)  # 经度
        # sheet1.write(i + 1, 2, g_addr_list[i], style)  # 公司地址
        # sheet1.write(i + 1, 3, r_name_list[i], style)  # 法定法人
        # sheet1.write(i + 1, 4, register_capital_list[i], style)  # 注册资本
        # sheet1.write(i + 1, 5, found_time_list[i], style)  # 成立日期
        # sheet1.write(i + 1, 6, g_status_list[i], style)  # 公司状态
        # sheet1.write(i + 1, 7, r_email_list[i], style)  # 法人邮箱
        # sheet1.write(i + 1, 8, r_phone_list[i], style)  # 法人电话
        # sheet1.write(i + 1, 9, taxpayer_code_list[i], style)  # 纳税人识别号
        # sheet1.write(i + 1, 10, register_institute_list[i], style)  # 登记机关
        # sheet1.write(i + 1, 11, register_code_list[i], style)  # 工商注册号
        # sheet1.write(i + 1, 12, g_desc_list[i], style)  # 公司简介
    # 保存excel文件，有同名的直接覆盖
    workbook.save(os.getcwd() + r"\\海岛数据.xls")
    print('保存完毕~')
